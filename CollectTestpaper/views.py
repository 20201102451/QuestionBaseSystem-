from django.shortcuts import render,HttpResponse
import docx
from django.db import transaction
from CollectTestpaper import models
# Create your views here.
def test(request):
    if request.method == 'GET':
        x = request.POST.get('lala');
        return render(request, 'test.html')
    else:
        global files;
        x_get = request.FILES.get('上传的文件')
        try:
            files = docx.Document(x_get)
        except Exception as e:
            return HttpResponse("插入失败：上传试卷类型有误！清仔细检查！")
        # testChoice()#获取选择题
        # testJudge()#获取判断题
        # testInput()#获取填空题
        # getTitle()#获取试卷名
        '''首先，获取试卷中的各项内容
            接下来，将各项内容录入数据库
            首先应该是试卷名的录入'''
        titleString = getTitle()
        choiceList = testChoice()
        judgeList = testJudge()
        inputList = testInput()
        Hard = request.POST.get("Hard")#读取试卷难度
        Apply_time = request.POST.get("Apply_time")#读取出卷时间
        Intro = request.POST.get("Intro")
        try:
            save_id_testpaper = transaction.savepoint()
            '''插入试卷'''
            if x_get == None:
                raise Exception("请先上传文件！")
            if models.Testpaper.objects.filter(name=titleString).exists():
                raise Exception("试卷已存在！")
            models.Testpaper.objects.create(name=titleString, hard=Hard, apply_time=Apply_time, intro=Intro)#Successful for test!

            '''接下来插入判断题'''
            for i in judgeList:
                for j in i.items():
                    models.Question.objects.create(stems=j[1][0], level=2)#Successful for test!
                    models.TrueOrFalse.objects.create(
                        question_id=models.Question.objects.filter(stems=j[1][0]).first(),
                        standard_answer=j[1][1],
                    )
                    models.TestpaperQuestion.objects.create(#插入试题_试卷对应表
                        testpaper_id=models.Testpaper.objects.filter(name=titleString).first(),
                        question_id=models.Question.objects.filter(stems=j[1][0]).first()
                    )

            '''接下来插入填空题'''
            for i in inputList:
                for j in i.items():
                    models.Question.objects.create(stems=j[1][0], level=3)#Successful for test!
                    models.Gapfill.objects.create(
                        question_id=models.Question.objects.filter(stems=j[1][0]).first(),
                        standard_answer=j[1][1],
                    )
                    models.TestpaperQuestion.objects.create(#插入试题_试卷对应表
                        testpaper_id=models.Testpaper.objects.filter(name=titleString).first(),
                        question_id=models.Question.objects.filter(stems=j[1][0]).first()
                    )
            '''接下来插入选择题'''
            for i in choiceList:
                for j in i.items():
                    models.Question.objects.create(stems=j[1][0], level=1)#插入试题表
                    models.Choice.objects.create(#选项插入
                        question_id=models.Question.objects.filter(stems=j[1][0]).first(),
                        choice1=j[1][2],
                        choice2=j[1][3],
                        choice3=j[1][4],
                        choice4=j[1][5],
                        standard_answer=j[1][1],
                    )
                    models.TestpaperQuestion.objects.create(#插入试题_试卷对应表
                        testpaper_id=models.Testpaper.objects.filter(name=titleString).first(),
                        question_id=models.Question.objects.filter(stems=j[1][0]).first()
                    )
            transaction.savepoint_commit(save_id_testpaper)
        except Exception as e:
            transaction.savepoint_rollback(save_id_testpaper)
            return HttpResponse("插入失败：" + str(e))
        return render(request, 'testSuccessful.html')
def testChoice():
    choiceList = []#选择题总容器
    k = 0
    line = 0
    op = 0
    for i in files.paragraphs:

        if i.style.name == "Heading 2":
            if i.text == "选择题":
                A = 0
                B = 0
                C = 0
                D = 0
                op = 0
                k = 1
                continue
            elif i.style.name == "Heading 2":
                break
        if op == 0:
            choiceDict = dict()  # 单个选择题容器字典
            question_answer = []  # 存储题干和标准答案
        if k == 1:
            t = i.text
            try:
                t.index("题干：")
                # print(line)
                # print(t[3:])#此为题干
                try:
                    rt = t.index("【")
                except Exception as e:
                    print("未识别到答案")
                else:
                    question_answer.append(t[3:rt])  # 将题干存储入结构
                    question_answer.append(t[rt+1:-1])#将答案存储入结构
                    op = 1
                    line += 1
                # question_answer.append(t[3:])  # 将题干存储入结构
                # op = 1
                # line += 1
            except Exception as e:#选择题中，识别不到题干二字的段落必然是选项，故异常处理需要将本次内容记录，并将下一次读入合并，再进行识别选项
                # print(t)
                if A == 0:
                    try:
                        a_index = t.index("A.")
                    except Exception as e:
                        A = 0
                        continue
                    else:
                        A = 1
                        try:
                            b_index = t.index("B.")
                        except Exception as e:
                            B = 0
                            str_A = t[a_index+2:]
                        else:
                            B = 1
                            str_A = t[a_index+2:b_index]
                            try:
                                c_index = t.index("C.")
                            except Exception as e:
                                C = 0
                                str_B = t[b_index+2:]
                            else:
                                C = 1
                                str_B = t[b_index+2:c_index]
                                try:
                                    d_index = t.index("D.")
                                except Exception as e:
                                    D = 0
                                    str_C = t[c_index+2:]
                                else:
                                    D = 1
                                    str_C = t[c_index+2:d_index]
                                    str_D = t[d_index+2:]
                                    # print(str_A + '\n')  # 此为选项A
                                    # print(str_B + '\n')  # 此为选项B
                                    # print(str_C + '\n')  # 此为选项C
                                    # print(str_D + '\n')  # 此为选项D
                                    question_answer.append(str_A)
                                    question_answer.append(str_B)
                                    question_answer.append(str_C)
                                    question_answer.append(str_D)
                                    choiceDict[line] = question_answer  # 存储入单个题目容器
                                    choiceList.append(choiceDict)  # 存储入总容器
                                    A = 0
                                    B = 0
                                    C = 0
                                    D = 0
                                    op = 0
                elif B == 0:
                    try:
                        b_index = t.index("B.")
                    except Exception as e:
                        B = 0
                    else:
                        B = 1
                        try:
                            c_index = t.index("C.")
                        except Exception as e:
                            C = 0
                            str_B = t[b_index+2:]
                        else:
                            C = 1
                            str_B = t[b_index+2:c_index]
                            try:
                                d_index = t.index("D.")
                            except Exception as e:
                                D = 0
                                str_C = t[c_index + 2:]
                            else:
                                D = 1
                                str_C = t[c_index + 2:d_index]
                                str_D = t[d_index + 2:]
                                # print(str_A + '\n')#此为选项A
                                # print(str_B + '\n')#此为选项B
                                # print(str_C + '\n')#此为选项C
                                # print(str_D + '\n')#此为选项D
                                question_answer.append(str_A)
                                question_answer.append(str_B)
                                question_answer.append(str_C)
                                question_answer.append(str_D)
                                choiceDict[line] = question_answer  # 存储入单个题目容器
                                choiceList.append(choiceDict)  # 存储入总容器
                                A = 0
                                B = 0
                                C = 0
                                D = 0
                                op = 0
                elif C == 0:
                    try:
                        c_index = t.index("C.")
                    except Exception as e:
                        C = 0
                    else:
                        C = 1
                        try:
                            d_index = t.index("D.")
                        except Exception as e:
                            D = 0
                            str_C = t[c_index + 2:]
                        else:
                            D = 1
                            str_C = t[c_index + 2:d_index]
                            str_D = t[d_index + 2:]
                            # print(str_A + '\n')  # 此为选项A
                            # print(str_B + '\n')  # 此为选项B
                            # print(str_C + '\n')  # 此为选项C
                            # print(str_D + '\n')  # 此为选项D
                            question_answer.append(str_A)
                            question_answer.append(str_B)
                            question_answer.append(str_C)
                            question_answer.append(str_D)
                            choiceDict[line] = question_answer  # 存储入单个题目容器
                            choiceList.append(choiceDict)  # 存储入总容器
                            A = 0
                            B = 0
                            C = 0
                            D = 0
                            op = 0
                elif D == 0:
                    try:
                        d_index = t.index("D.")
                    except Exception as e:
                        D = 0
                    else:
                        D = 1
                        str_D = t[d_index+2:]
                        # print(str_A + '\n')  # 此为选项A
                        # print(str_B + '\n')  # 此为选项B
                        # print(str_C + '\n')  # 此为选项C
                        # print(str_D + '\n')  # 此为选项D
                        question_answer.append(str_A)
                        question_answer.append(str_B)
                        question_answer.append(str_C)
                        question_answer.append(str_D)
                        choiceDict[line] = question_answer  # 存储入单个题目容器
                        choiceList.append(choiceDict)  # 存储入总容器
                        A = 0
                        B = 0
                        C = 0
                        D = 0
                        op = 0
    return choiceList

def testJudge():
    judgeList = []#判断题总容器
    k = 0
    line = 1
    for i in files.paragraphs:
        judgeDict = dict()  # 单个判断题容器字典
        question_answer = []  # 存储题干和标准答案
        if i.style.name == "Heading 2":
            if i.text == "判断题":
                k = 1
                continue
            elif i.text == "填空题":
                break;
        if k == 1:
            t = i.text
            try:
                t.index("题干：")
            except Exception as e:#空白行，不予读入
                continue
            else:
                try:
                    rt = t.index("【")
                except Exception as e:
                    print("未识别到答案")
                else:
                    question_answer.append(t[3:rt])  # 将题干存储入结构
                    question_answer.append(t[rt+1:-1])#将答案存储入结构
                    judgeDict[line] = question_answer  # 存储入单个题目容器
                    judgeList.append(judgeDict)  # 存储入总容器
                    line += 1
    return judgeList

def testInput():
    #choiceList= [{"1":["题干","选项1","选项2","选项3","选项4","标准答案"]},]
    #judgeList= [{"1":["题干","答案"]},]
    #inputList= [{"1":["题干","答案"]},]
    k = 0
    line = 1
    inputList = []#填空题总容器

    for i in files.paragraphs:
        inputDict = dict()  # 单个填空题容器字典
        question_answer = []  # 存储题干和标准答案
        if i.style.name == "Heading 2":
            if i.text == "填空题":
                k = 1
                continue
        if k == 1:
            t = i.text
            try:
                t.index("题干：")
            except Exception as e:#空白行，不予读入
                continue
            else:
                try:
                    rt = t.index("【")
                except Exception as e:
                    print("未识别到答案")
                else:
                    question_answer.append(t[3:rt])  # 将题干存储入结构
                    question_answer.append(t[rt+1:-1])#将答案存储入结构
                    inputDict[line] = question_answer  # 存储入单个题目容器
                    inputList.append(inputDict)  # 存储入总容器
                    line += 1
    return inputList

def getTitle():#获取试卷名
    for i in files.paragraphs:
        if i.style.name == "Heading 1":
            return i.text
